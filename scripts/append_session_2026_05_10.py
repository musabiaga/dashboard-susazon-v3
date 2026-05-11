#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Agrega un apéndice "Sesión 2026-05-10" a los 6 .docx ya generados,
con resumen de Mejoras 1-7 + Branding InCom + Login rediseño.
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DOCS = Path("/Users/mauusabiaga/Downloads/DASHBOARD SEMANAL VENTAS V3.0 [Claude Code]/docs")

# ============================================================
# Contenido del apéndice por documento (cada doc recibe lo
# que le aplica a su scope).
# ============================================================

APPENDIX_COMMON_HEADER = "Sesión 2026-05-10 — Mejoras 1-7 + Branding InCom + Login premium"

APPENDIX_ARQUITECTURA = """
## Cambios arquitectónicos (sesión 2026-05-10)

### Nuevo helper: lib/aggregate.ts (Mejora 7)
Implementa agregación dinámica multi-territorio en cliente.
Funciones puras:
- aggregateKpis(kpis): suma N TerritoryKpi recomputando marginPct desde sumas.
- aggregateDimensionRows(rowsList): suma DimensionRow[] por nombre (18 campos numéricos).
- aggregatePerdidoRows(rowsList): suma PerdidoRow[] por no_cliente (27 campos).
- aggregateBudget(territories, selected): suma de ventaBudget.
- selectedKpis() / rowsBySelected(): helpers de mapeo Set → arrays.

Performance: O(N × items_por_territorio). Memoizado con useMemo en DashboardClient.

### 4 modos de selección de territorio (Mejora 7)
DashboardClient ahora maneja:
- single: 1 territorio individual seleccionado (uni-select sidebar)
- aggregated-all: "Todos" + set completo de activos → usa pre-agregado del server (rápido)
- aggregated-custom: "Todos" + set personalizado (subset configurable con ⚙️)
- aggregated-none: "Todos" + set vacío → empty state

### Nuevo helper: lib/export-excel.ts (Mejora 6)
Generador Excel reusable con exceljs (no xlsx por CVEs activas).
Estructura por hoja: bloque resumen → header bold + freeze panes → filas con zebra → row TOTAL.
numFmt nativo de Excel (los % son proporciones reales 0.0%, pivot-ready).
Lazy import → exceljs ~700KB solo se carga al click.

### Nuevo endpoint: /api/dashboard/clientes-dia (Mejora 1)
GET con query params year/month/day/territorio (opcional).
Query a kpi_cliente_diario (RLS via security_invoker).
Agrega por no_cliente, sort por venta desc.
Cache HTTP: private, max-age=60s (días históricos son inmutables).

### Nuevas migraciones SQL
- 011_kpi_diario_dimensiones.sql: 4 vistas diarias (grupo, sku, cliente, vendedor)
  para día-vs-día YoY (Mejora 2). Todas con security_invoker=true.
- 012_kpi_cliente_perdidos_margen.sql: agrega mes_margen + ytd_margen al final de
  kpi_cliente_perdidos (CREATE OR REPLACE solo permite append al final).

### Helper findCalendarDayForBizDays() en lib/business-days.ts
Mapea día hábil N de un año a su día calendario equivalente en otro año
(considerando feriados LFT y domingos distintos). Crítico para comparativos
día-vs-día precisos entre años.
"""

APPENDIX_DICCIONARIO = """
## Nuevas estructuras de datos (sesión 2026-05-10)

### DimensionRow (extendido con campos al-día)
Cada fila ahora incluye 18 campos numéricos: cierre + al-día × venta/kg/margen × 3 años.
Permite comparativos día-vs-día equitativos en charts y tablas.

### PerdidoRow (extendido con margen)
Ahora incluye mes_margen_* y ytd_margen_* (3 años × mes/ytd × al-día/cierre).
27 campos numéricos totales por cliente.

### ClienteDelDia (nuevo)
Estructura del endpoint /api/dashboard/clientes-dia:
- no_cliente: código ERP
- cliente: razón social
- vendedor: nombre asignado
- venta, margen, kg, marginPct

### Quote (descartado en revisión final)
Estructura propuesta en lib/login-quotes.ts para frases rotativas de pensadores.
Implementado y luego ELIMINADO por feedback del usuario (Mauricio prefirió logo sin frases).
Conservado en el git history (commit 3ade72e) por si se quiere retomar.

### Set<string> (selectedTerritories vs aggregatedTerritories)
DashboardClient tiene 2 sets distintos:
- selectedTerritory: string (uni-select, no persiste)
- aggregatedTerritories: Set<string> (qué incluye "Todos", persiste en localStorage
  con key "dashboard-aggregated-territories")

### ExcelColumn, ExcelSummaryRow, ExcelExportOptions
Tipos del helper lib/export-excel.ts.
- ExcelColumn: { header, key, width?, numFmt?, align? }
- ExcelSummaryRow: { label, value, numFmt? }
- ExcelExportOptions: { fileName, sheetName, title?, subtitle?, summary?, columns, rows, totalRow? }

### Nuevos archivos de assets
- public/incom-logo.png (1254×1254 con fondo papel — source)
- public/incom-mark.png (256×256 transparente)
- public/incom-mark@2x.png (512×512 retina)
- public/incom-mark@4x.png (1024×1024 hero login)
- app/icon.png (48×48 favicon)
- app/apple-icon.png (180×180)
- app/opengraph-image.png (1200×630)
- app/twitter-image.png (copia del OG)
"""

APPENDIX_CHANGELOG = """
## v3.3.0 — 2026-05-10 · Branding InCom + Login premium

### Added
- Branding InCom completo (favicon + apple-icon + OG image + Twitter card + metadata).
- Login rediseñado split-screen estilo enterprise (hero con escudo InCom + aurora animado + logo Susazón auto-switch).
- Intro animation 1.2s con prefers-reduced-motion respetado.
- Helper lib/aggregate.ts con 6 funciones puras para multi-territorio.
- Helper lib/export-excel.ts con generador Excel reusable.
- Endpoint /api/dashboard/clientes-dia para lazy load Tracking Diario.
- Componente ExportExcelButton.tsx en los 7 tabs.

### Changed
- Sidebar: item "Todos" con ícono ⚙️ Settings2 que abre dropdown configurable.
- PerdidosTab: eliminado TerritoryFilter local (~150 líneas) → sidebar global.
- DashboardClient: 4 modos de selección (single/aggregated-all/custom/none).
- Tema Liquid Glass: clase .frost-popover con bg opaco + blur 40px.
- SusazonLogo: nueva prop surface="header"|"page" con PAGE_DARK_THEMES.

### Fixed
- Día-vs-día YoY con precisión 100% (Mejora 2: 4 vistas SQL + findCalendarDayForBizDays).
- Bug Perdidos: comparación mes-cierre vs mes-actual inflaba "declive".
- Vendedores: cambiar Sus/Suve preservaba selección inválida (key remount fix).
- Theme dark: logo Susazón se veía marrón ilegible (PAGE_DARK_THEMES expandido).

### Migrations
- 011_kpi_diario_dimensiones.sql
- 012_kpi_cliente_perdidos_margen.sql

## v3.2.0 — 2026-05-09 · Mejoras 1-5

- Mejora 1: Lazy endpoint clientes/día Tracking Diario.
- Mejora 2: Día-vs-día YoY (Commits A datos, B charts, C tablas).
- Mejora 3: Multi-select Productos persistente + lock 15.
- Mejora 4: Multi-select Clientes con MultiSelectChips.
- Mejora 5: Columnas KG en tablas inferiores.
- PerdidosTab: chips status + dona + LossCards + buscador amplio + año 2024.

## v3.1.0 — 2026-05-01 · Fase 2d completa

- 7 tabs funcionando con data real.
- 16 territorios + PTTO 2026.
- 14 usuarios autorizados de 15 planeados.
- Admin Panel completo.

## v3.0.0 — 2026-04-26 · Lanzamiento V3.0

- Next.js 16 + Supabase + Vercel.
- Auth con RLS por territorio.
- Dominio custom dashboardcomercialsusazon.com.
"""

APPENDIX_MANUAL = """
## Nuevas features para el usuario (sesión 2026-05-10)

### Multi-select global de territorios ("Todos" configurable)
En el sidebar, al lado del item "Todos" verás un ícono ⚙️ pequeño.
- Click en "Todos" → muestra el agregado de todos los territorios activos.
- Click en el ⚙️ → abre un dropdown con checkboxes para elegir QUÉ territorios incluye "Todos".
- Por ejemplo: marcar solo Mérida + Cancún + CDMX → click en "Todos" → ves solo esos 3 sumados.
- La configuración persiste entre sesiones (localStorage).
- Si seleccionas un territorio individual, la config de "Todos" NO afecta.

### Export Excel en los 7 tabs
Cada tab tiene un botón "Exportar Excel" arriba a la derecha.
- Respeta los filtros activos en pantalla (WYSIWYG).
- Incluye bloque resumen + tabla completa + row TOTAL.
- Los porcentajes son números reales (pivot-ready en tablas dinámicas).
- Naming: [Tab]_[Territorio]_2026-05-10.xlsx.

### Tab Tracking Diario: expandir clientes por día
Cada fila de la tabla diaria tiene una flecha (▶) que despliega los clientes del día.
- Click → fetch lazy desde Supabase, aparece sub-tabla.
- Muestra: Cliente, Vendedor, Venta $, % del día, KG, Margen $, Margen %.
- Cache en memoria: re-expandir un día es instantáneo.

### Día-vs-día YoY (precisión 100%)
Los charts y tablas ahora comparan "mismo día laboral" entre años.
- Antes: comparaba mes-cierre 2025 vs mes-actual 2026 (parcial) → sesgo.
- Ahora: compara hasta el mismo día hábil considerando feriados LFT distintos.
- "Var % vs 25" es preciso al día.

### Login rediseñado
Layout split-screen con escudo InCom grande a la izquierda y form a la derecha.
- Intro animada de 1.2 segundos al cargar.
- Logo Susazón Gourmet cambia color automáticamente según el theme.
- 6 themes disponibles incluyendo Liquid Glass (aurora animado de Apple).
"""

APPENDIX_GUIA_TI = """
## Cambios TI relevantes (sesión 2026-05-10)

### Nueva dependencia: exceljs
Agregado para export Excel en los 7 tabs.
Lazy import → solo se descarga cuando el usuario hace click en "Exportar Excel".
Bundle inicial sin impacto.
Reemplaza a la opción xlsx que tenía CVE-2023-30533 (prototype pollution) y
CVE-2024-22363 (ReDoS) sin fix activo.

### Nuevas migraciones SQL aplicadas en producción
- 011_kpi_diario_dimensiones.sql (4 vistas diarias)
- 012_kpi_cliente_perdidos_margen.sql (agrega margen)

Aplicación: vía Supabase MCP o SQL Editor del Dashboard.
Idempotentes: CREATE OR REPLACE VIEW.

### Nuevos endpoints API
- GET /api/dashboard/clientes-dia?year=&month=&day=&territorio=
  Auth requerida. RLS aplicado via security_invoker.
  Cache: private, max-age=60s.

### Nuevos assets en /public y /app
- public/incom-logo.png, incom-mark.png, incom-mark@2x.png, incom-mark@4x.png
- app/icon.png, app/apple-icon.png, app/opengraph-image.png, app/twitter-image.png
- public/susazon-logo.png + susazon-logo-light.png (sin cambios)

### Variables de entorno (sin cambios)
Las mismas que antes. Verificar contra .env.example.

### Persistencia adicional en localStorage (cliente)
- dashboard-aggregated-territories: Set<string> de territorios incluidos en "Todos".
- perdidos-metric-mode, perdidos-status-filter: configuración del tab Perdidos.
- productos-selected-skus, clientes-selected, vendedores-selected: multi-selects.
- tracking-diario-mode: pesos vs kilos.

### Verificación post-deploy recomendada
1. Verificar 6 themes (clean, editorial, warm-neo, supabase-orange, stock-market, liquid-glass).
2. Verificar export Excel en los 7 tabs.
3. Verificar multi-select global del sidebar (single / aggregated-all / custom / none).
4. Verificar Tracking Diario expand (debe cargar clientes del día en <500ms).
5. Verificar OG image al pegar URL en WhatsApp/Slack/iMessage.
"""

APPENDIX_RECONSTRUCCION = """
## Adiciones para reconstrucción (sesión 2026-05-10)

### Paso 11 — Helpers de agregación (nuevo)
Crear lib/aggregate.ts con 6 funciones puras:
1. emptyKpi() → TerritoryKpi vacío
2. aggregateKpis(kpis[]) → suma con recomputo de marginPct
3. aggregateDimensionRows(rowsList[][]) → suma por nombre
4. aggregatePerdidoRows(rowsList[][]) → suma por no_cliente
5. aggregateBudget(territories, selected) → suma de ventaBudgets
6. selectedKpis() / rowsBySelected() → helpers de mapeo

Ver código completo en /lib/aggregate.ts del repo.

### Paso 12 — Helper de Excel (nuevo)
Crear lib/export-excel.ts con exceljs.
API: exportToExcel({ fileName, sheetName, title, subtitle, summary, columns, rows, totalRow }).
- Header con freeze panes en la primer fila de tabla.
- Zebra rows.
- numFmt nativo Excel ($#,##0, 0.0%, #,##0).
- Lazy import.

### Paso 13 — Endpoint lazy clientes/día (nuevo)
Crear app/api/dashboard/clientes-dia/route.ts:
- GET con validación rangos (year >= 2024, month 1-12, day 1-31).
- Auth check + RLS via security_invoker.
- Query a kpi_cliente_diario filtrada por territorio (opcional).
- Agregación por no_cliente.
- Cache: private, max-age=60s.

### Paso 14 — Sidebar con item "Todos" + ⚙️ configurable (nuevo)
Refactor components/dashboard/Sidebar.tsx:
- Props nuevas: aggregatedTerritories: Set<string>, onToggleAggregated, onToggleAllAggregated.
- Componente AggregatedItem: botón principal + ícono Settings2 separado.
- Dropdown con clase frost-popover (override CSS para liquid-glass en globals.css).
- Estado abre/cierra con click outside detection.

### Paso 15 — DashboardClient 4 modos (nuevo)
Refactor app/dashboard/DashboardClient.tsx:
- Estado: selectedTerritory: string + aggregatedTerritories: Set<string>.
- 4 modos derivados: single / aggregated-all / aggregated-custom / aggregated-none.
- Helpers resolveDimRows() y resolvePerdidoRows() que llaman aggregate*.
- Persistencia localStorage solo para aggregatedTerritories.
- Auto-add nuevos territorios reactivados al set.

### Paso 16 — Login rediseñado (nuevo)
Reescribir app/login/page.tsx con layout split-screen:
- Hero izquierdo (60%): escudo InCom + subtítulo + aurora animado + grid overlay.
- Card derecho (40%): logo Susazón 6x + form + footer copyright.
- Intro 1.2s con CSS animations + prefers-reduced-motion.
- Logo Susazón con prop surface="page" (auto-switch por theme).

### Paso 17 — Thumbnails de branding (nuevo)
Generar con Pillow desde public/incom-logo.png:
1. Combinar 2 PNGs del SVG (máscara grayscale + escudo RGB) para alpha exacto.
2. Crop al bbox del contenido.
3. Padding cuadrado transparente.
4. Resize a 48, 180, 256, 512, 1024 px.
5. app/icon.png + apple-icon.png + opengraph-image.png + twitter-image.png.
6. Actualizar app/layout.tsx con metadataBase + openGraph + twitter cards.

Ver script Python completo en scripts/gen_docs.py (mencionado en commits 977ddc3 y 0f5572e).

### Verificación post-reconstrucción (12 puntos)
1. Aurora del login se mueve (animation loop 18s).
2. Intro corre 1.2s al cargar el login.
3. Logo InCom 520px visible en hero izquierdo.
4. Logo Susazón 6x en card derecho.
5. Theme switch funciona en los 6 themes.
6. Item "Todos" del sidebar con ⚙️ visible.
7. Dropdown del ⚙️ funciona y persiste.
8. Export Excel funciona en los 7 tabs.
9. Tracking Diario expand muestra clientes/día.
10. Charts día-vs-día (no mes-completo).
11. Perdidos SIN TerritoryFilter local (un solo place).
12. OG image al compartir URL en WhatsApp.
"""

CONTENT_BY_DOC = {
    "01_Arquitectura_Tecnica.docx": APPENDIX_ARQUITECTURA,
    "02_Diccionario_Datos.docx": APPENDIX_DICCIONARIO,
    "03_ChangeLog_Release_Notes.docx": APPENDIX_CHANGELOG,
    "04_Manual_Usuario.docx": APPENDIX_MANUAL,
    "05_Guia_TI_Despliegue.docx": APPENDIX_GUIA_TI,
    "06_Guia_Reconstruccion.docx": APPENDIX_RECONSTRUCCION,
}


def add_appendix(doc_path: Path, body: str):
    doc = Document(str(doc_path))

    # Page break antes del apéndice
    doc.add_page_break()

    # Título del apéndice (H1)
    h1 = doc.add_heading("Apéndice — Sesión 2026-05-10", level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Sub-título
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Mejoras 1-7 · Branding InCom · Login premium")
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("")  # spacer

    # Body: parsear líneas con ## como H2, ### como H3, otras como párrafo
    for line in body.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("  "):
            # Sub-item indentado
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.left_indent = 360_000  # 0.25"
        else:
            doc.add_paragraph(line)

    doc.save(str(doc_path))


print("Agregando apéndices de sesión 2026-05-10 a los 6 .docx...\n")
for fname, body in CONTENT_BY_DOC.items():
    path = DOCS / fname
    if not path.exists():
        print(f"  ✗ {fname} NO existe — skipped")
        continue
    add_appendix(path, body)
    size_kb = path.stat().st_size // 1024
    print(f"  ✓ {fname} ({size_kb} KB)")

print("\n✅ Apéndices agregados.")
