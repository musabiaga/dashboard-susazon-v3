#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de los 6 documentos Word del Dashboard Comercial Susazón V3.0.

Uso:
    python3 scripts/gen_docs.py

Salida en: docs/
    - 01_Arquitectura_Tecnica.docx
    - 02_Diccionario_Datos.docx
    - 03_ChangeLog_Release_Notes.docx
    - 04_Manual_Usuario.docx
    - 05_Guia_TI_Despliegue.docx
    - 06_Guia_Reconstruccion.docx
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# Constantes del proyecto
# ============================================================
PROYECTO = "Dashboard Comercial Susazón V3.0"
EMPRESA = "Grupo Susazón (Susazón + Suve)"
OWNER = "Mauricio Usabiaga, Director de Operaciones"
VERSION = "1.0.0"
FECHA = date.today().strftime("%Y-%m-%d")
REPO = "github.com/musabiaga/dashboard-susazon-v3"
URL_PROD = "dashboard-susazon-v3-44sp-hw6gg0rwb-musabiagas-projects.vercel.app"

DOCS_DIR = Path(__file__).resolve().parent.parent / "docs"
DOCS_DIR.mkdir(exist_ok=True)


# ============================================================
# Helpers de formato
# ============================================================
def add_cover(doc: Document, titulo: str, subtitulo: str):
    """Portada estandarizada."""
    # Espacio arriba
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROYECTO)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xED, 0x68, 0x08)  # Naranja Susazón
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.font.size = Pt(28)
    run.bold = True

    if subtitulo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitulo)
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Espacio
    for _ in range(8):
        doc.add_paragraph()

    # Metadata box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Versión {VERSION} · {FECHA}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{EMPRESA}\n{OWNER}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()


def add_h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xED, 0x68, 0x08)


def add_h2(doc: Document, text: str):
    doc.add_heading(text, level=2)


def add_h3(doc: Document, text: str):
    doc.add_heading(text, level=3)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_code(doc: Document, code: str):
    """Bloque de código con font monospace y bg gris claro."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # Para que se vea más como código, fondo gris claro con shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    return p


def add_table(doc: Document, headers: list, rows: list[list]):
    """Tabla con header gris y rows alternados."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        # Background naranja Susazón claro
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "FFE5D0")
        tcPr.append(shd)

    # Rows
    for ridx, row in enumerate(rows):
        cells = table.rows[ridx + 1].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)

    return table


def add_footer(doc: Document, text: str):
    """Footer en cada página."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ============================================================
# DOC 01 — Arquitectura Técnica
# ============================================================
def gen_arquitectura():
    doc = Document()

    add_cover(doc, "Arquitectura Técnica", "Diseño del sistema, stack y flujo de información")

    add_h1(doc, "Resumen Ejecutivo")
    add_para(
        doc,
        f"{PROYECTO} es una aplicación web SaaS interna que reemplaza al Dashboard V2.2 (single-page HTML "
        "alojado en Netlify). El objetivo es proveer a 15 usuarios internos de Grupo Susazón una herramienta "
        "comercial moderna con permisos granulares, datos en tiempo casi-real desde APIs de Susazón y Suve, "
        "y panel administrativo completo — sin exponer credenciales sensibles al navegador del usuario."
    )
    add_para(
        doc,
        "El V2.2 tenía 3 problemas críticos que motivaron esta reescritura: (1) la API Key de Susazón estaba "
        "hardcoded en el HTML del cliente y era visible desde DevTools; (2) no había sistema de permisos por "
        "usuario — todos los autenticados veían todos los territorios; (3) bugs visuales pendientes en charts "
        "(eje X mal renderizado). El V3.0 resuelve los tres."
    )

    add_h1(doc, "Stack Tecnológico")
    add_table(
        doc,
        ["Capa", "Tecnología", "Versión", "Justificación"],
        [
            ["Framework", "Next.js (App Router + Turbopack)", "16.2.4", "Server Components esconden secretos"],
            ["UI", "React + TypeScript + Tailwind CSS", "19.2 / 5 / 4", "Modern, type-safe, CSS-first"],
            ["Charts", "Recharts", "3.8", "Native React, no Chart.js"],
            ["DB + Auth", "Supabase Postgres", "Free tier", "RLS nativo, auth incluido"],
            ["Auth SDK", "@supabase/ssr", "0.10", "SSR-aware con cookies"],
            ["Hosting", "Vercel", "Hobby plan", "Deploy desde GitHub, edge network"],
            ["Repo", "GitHub", "privado", "Backup + versionado"],
            ["Icons", "lucide-react", "0.4+", "Consistente, tree-shakeable"],
            ["Validación", "Zod", "3.x", "Runtime type-safety"],
        ],
    )

    add_h1(doc, "Arquitectura de Alto Nivel")
    add_para(doc, "Diagrama de flujo de información (descripción textual):")
    add_code(
        doc,
        """
Browser (usuario autenticado)
   │
   │  HTTP request con cookie de sesión Supabase
   │
   ▼
Next.js Edge (Vercel)
   │
   ├─→ proxy.ts (middleware): valida sesión, redirect si no auth
   │
   ├─→ Server Component renderiza con datos de Supabase (RLS aplicado)
   │
   └─→ API Route handler (server-side):
            │
            ├─→ Lee process.env.SUSAZON_API_KEY (servidor only)
            ├─→ POST a https://sasweb.susazon.mx/.../api_ERPPyMEDashboard/
            ├─→ Recibe filas de venta (~5,000/mes Susazón, ~2,500/mes Suve)
            ├─→ filterValidRows: descarta kg = 0
            ├─→ INSERT en sales_rows (idempotente: DELETE before INSERT)
            └─→ Audit log entry
   ▼
Supabase Postgres
   │
   ├─→ RLS: SELECT filtrado por visible_territories_for_current_user()
   ├─→ Vistas KPI con security_invoker = true (heredan RLS)
   └─→ Tablas core: sales_rows, users_permissions, territories_state,
                    territory_budgets, audit_log
        """,
    )

    add_h1(doc, "Componentes y Carpetas")
    add_table(
        doc,
        ["Carpeta", "Propósito"],
        [
            ["app/", "Páginas Next.js (App Router)"],
            ["app/dashboard/", "Dashboard principal con 7 tabs"],
            ["app/admin/", "Panel admin (territorios, usuarios, audit)"],
            ["app/cargar-datos/", "Refresh APIs + editor PTTO"],
            ["app/api/", "API routes server-side (proxy seguro)"],
            ["components/dashboard/", "Componentes de los 7 tabs y charts"],
            ["components/theme/", "6 themes + selector modal"],
            ["components/layout/", "Header, layout shells"],
            ["components/brand/", "Logo Susazón con auto-switch"],
            ["lib/supabase/", "Clientes Supabase (client/server/admin)"],
            ["lib/susazon-api.ts", "Wrapper server-side de APIs Susazón/Suve"],
            ["lib/format.ts", "Formatters (money, kilos, dates)"],
            ["lib/business-days.ts", "Cálculo días hábiles L-S menos LFT"],
            ["lib/admin-guards.ts", "requireAdmin() helper"],
            ["lib/themes.ts", "Definición de los 6 themes"],
            ["supabase/migrations/", "10 migraciones SQL aplicadas"],
            ["docs/", "Esta documentación"],
            ["proxy.ts", "Middleware Next.js 16 (renombrado)"],
            [".env.local", "Secrets (NO commit)"],
            ["AGENTS.md", "Auto-cargado para Claude. Contexto crítico"],
        ],
    )

    add_h1(doc, "Flujo de Datos End-to-End")
    add_h2(doc, "Refresh APIs (admin/director)")
    add_para(
        doc,
        "El usuario admin/director va a /cargar-datos, selecciona rango de meses y fuentes (Susazón/Suve), "
        "y dispara el refresh. La API route /api/data/refresh:"
    )
    add_bullet(doc, "Verifica auth + rol (admin o director)")
    add_bullet(doc, "Lee SUSAZON_API_KEY y SUSAZON_API_URL del proceso (NUNCA del cliente)")
    add_bullet(doc, "Por cada (source, mes), POST a la API → recibe rows")
    add_bullet(doc, "Borra (source, año, mes) existentes en sales_rows (idempotencia)")
    add_bullet(doc, "INSERT bulk de las nuevas rows con kg > 0")
    add_bullet(doc, "Inserta evento data_refresh en audit_log")
    add_bullet(doc, "Retorna { sources_processed, rows_imported, errors[] }")

    add_h2(doc, "Visualización (cualquier usuario autenticado)")
    add_para(
        doc,
        "Cuando el user accede a /dashboard, el Server Component ejecuta queries en paralelo a las vistas KPI. "
        "Cada query trae solo filas de territorios permitidos para ese usuario, gracias a la combinación de:"
    )
    add_bullet(doc, "RLS policy en sales_rows: territorio = ANY(visible_territories_for_current_user())")
    add_bullet(doc, "Views KPI con WITH (security_invoker = true) — heredan RLS al consultar")
    add_bullet(doc, "Helper visible_territories_for_current_user() retorna intersección de allowed_territories × is_active=true")

    add_h1(doc, "Seguridad")
    add_h2(doc, "Capas de defensa")
    add_table(
        doc,
        ["Capa", "Mecanismo", "Protege contra"],
        [
            ["1. Browser", "JWT cookie httpOnly, anon key con permisos limitados", "Robo de cookie via XSS"],
            ["2. Edge / Middleware", "proxy.ts valida sesión antes de servir páginas privadas", "Acceso sin auth"],
            ["3. API Routes", "Validan rol + permisos antes de cualquier op", "Privilege escalation"],
            ["4. Database", "RLS policies + security_invoker views", "SQL injection / API directo"],
            ["5. Audit Log", "Tabla inmutable con todos los eventos críticos", "Repudio / forensics"],
        ],
    )

    add_h2(doc, "Secretos NUNCA expuestos al cliente")
    add_bullet(doc, "SUSAZON_API_KEY — solo en process.env del server")
    add_bullet(doc, "SUVE_API_KEY — solo en process.env del server")
    add_bullet(doc, "SUPABASE_SERVICE_ROLE_KEY — solo en API routes con permisos admin")
    add_para(
        doc,
        "Verificación QA: se escanearon 24 archivos JS/CSS del bundle de producción y NINGUNO contenía las "
        "claves prohibidas. Solo el NEXT_PUBLIC_SUPABASE_ANON_KEY aparece (esperado, es la anon key con RLS aplicada)."
    )

    add_h1(doc, "Decisiones Arquitectónicas (ADRs)")
    add_para(doc, "Las decisiones críticas están documentadas en SESSION_LOG.md como D001-D008. Resumen:")
    add_bullet(doc, "D001: Reescritura completa Next.js + Supabase (justificación: el V2.2 es inseguro)")
    add_bullet(doc, "D002: RLS a nivel DB, no solo UI (defensa en profundidad)")
    add_bullet(doc, "D003: Mapeo PTTO Intercompañias y Zurt-t (no renombrar nombres del API)")
    add_bullet(doc, "D004: Bumpear Supabase max-rows a 50000 (vs crear vistas no-territoriadas)")
    add_bullet(doc, "D005: 6 themes con selector modal (UX > dropdown apretado)")
    add_bullet(doc, "D006: Un solo proyecto Supabase para esta escala (15 usuarios, single-tenant)")
    add_bullet(doc, "D007: Vercel Hobby plan + maxDuration 300s (alcanza 95% del uso real)")
    add_bullet(doc, "D008: Safari fix Liquid Glass (isolation + transform translateZ)")

    add_h1(doc, "Limitaciones Conocidas")
    add_bullet(doc, "Hobby plan Vercel limita refresh a ~5 meses con ambas APIs. Solución: rangos chicos o upgrade Pro.")
    add_bullet(doc, "No hay staging environment — deploys van directo a prod desde main branch.")
    add_bullet(doc, "Audit log no tiene retention policy automática (crece indefinidamente).")
    add_bullet(doc, "Single project Supabase — no hay sandbox para experimentar sin afectar prod.")
    add_bullet(doc, "Suve API es lenta (~60s/mes) — diseño de SQL Express, no podemos cambiar.")

    add_footer(doc, f"{PROYECTO} · Arquitectura Técnica · {VERSION}")
    out = DOCS_DIR / "01_Arquitectura_Tecnica.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")


# ============================================================
# DOC 02 — Diccionario de Datos
# ============================================================
def gen_diccionario():
    doc = Document()
    add_cover(doc, "Diccionario de Datos", "Schemas de DB, contratos de APIs y estructuras internas")

    add_h1(doc, "Tablas de Postgres")

    add_h2(doc, "sales_rows (tabla principal)")
    add_para(doc, "Persiste todas las filas de venta importadas desde Susazón y Suve. RLS habilitado.")
    add_table(
        doc,
        ["Columna", "Tipo", "Notas"],
        [
            ["id", "bigserial", "PK auto-incremental"],
            ["empresa", "smallint CHECK (empresa IN (0,1))", "0=Susazón, 1=Suve"],
            ["no_cliente", "text", "ID del cliente desde el ERP"],
            ["cliente", "text", "Nombre del cliente"],
            ["territorio", "text", "Nombre del territorio (filtrado por RLS)"],
            ["vendedor", "text", "Nombre del vendedor"],
            ["sku", "text", "Código del producto"],
            ["kg", "numeric", "Solo se insertan filas con kg > 0"],
            ["fecha", "date", "Fecha de la venta YYYY-MM-DD"],
            ["anio", "smallint", "Año extraído de fecha"],
            ["mes", "smallint", "Mes 1-12"],
            ["venta", "numeric", "Monto en pesos MXN"],
            ["margen", "numeric", "Margen en pesos MXN"],
            ["familia", "text", "Familia de producto (legacy)"],
            ["grupo", "text", "Grupo de producto (campo nuevo agregado por TI)"],
            ["created_at", "timestamptz", "Default now()"],
        ],
    )
    add_para(doc, "RLS Policy:")
    add_code(doc, 'CREATE POLICY "users_read_visible_territories" ON sales_rows\n'
                  '  FOR SELECT USING (territorio = ANY (visible_territories_for_current_user()));')

    add_h2(doc, "users_permissions")
    add_para(doc, "Extiende auth.users con rol y permisos por territorio.")
    add_table(
        doc,
        ["Columna", "Tipo", "Notas"],
        [
            ["user_id", "uuid PK", "FK a auth.users(id) ON DELETE CASCADE"],
            ["email", "text UNIQUE", ""],
            ["full_name", "text", ""],
            ["role", "user_role enum", "admin | director | gerente_regional | vendedor"],
            ["allowed_territories", "text[]", "NULL = todos. Array vacío = ninguno"],
            ["can_edit_ptto", "boolean", "Default false. True para admin/director"],
            ["theme_preference", "theme_preference enum", "clean | editorial | warm-neo (default clean)"],
            ["is_active", "boolean", "False = soft delete (no elimina, desactiva login)"],
            ["created_at", "timestamptz", ""],
            ["updated_at", "timestamptz", "Trigger set_updated_at() lo mantiene"],
            ["last_login", "timestamptz", "Actualizado por callback de auth"],
        ],
    )

    add_h2(doc, "territories_state")
    add_para(doc, "Toggle global on/off por territorio. Auto-poblado por trigger cuando llega un territorio nuevo en sales_rows.")
    add_table(
        doc,
        ["Columna", "Tipo", "Notas"],
        [
            ["territory_name", "text PK", ""],
            ["is_active", "boolean", "False = apagado para todos"],
            ["disabled_by", "uuid", "FK a auth.users(id)"],
            ["disabled_at", "timestamptz", ""],
            ["reason", "text", "Motivo opcional (max 500 chars)"],
            ["updated_at", "timestamptz", ""],
        ],
    )

    add_h2(doc, "territory_budgets")
    add_para(doc, "PTTO mensual por territorio. Editable solo por admin/director.")
    add_table(
        doc,
        ["Columna", "Tipo", "Notas"],
        [
            ["territorio", "text", "Parte del PK compuesto"],
            ["anio", "smallint", "Parte del PK compuesto"],
            ["mes", "smallint", "Parte del PK compuesto"],
            ["venta_budget", "numeric", "Pesos MXN"],
            ["updated_by", "uuid", "FK a auth.users(id)"],
            ["updated_at", "timestamptz", ""],
        ],
    )

    add_h2(doc, "audit_log (inmutable)")
    add_para(doc, "Registro de eventos de seguridad y cambios sensibles. Sin políticas UPDATE/DELETE — inmutable.")
    add_table(
        doc,
        ["Columna", "Tipo", "Notas"],
        [
            ["id", "uuid PK", "gen_random_uuid()"],
            ["user_id", "uuid", "FK a auth.users(id) ON DELETE SET NULL"],
            ["user_email", "text", "Snapshot del email al momento del evento"],
            ["action", "audit_action enum", "Ver enum abajo"],
            ["details", "jsonb", "Contexto: before/after, changed_fields, etc."],
            ["ip_address", "inet", "Opcional"],
            ["user_agent", "text", "Opcional"],
            ["created_at", "timestamptz", "Default now()"],
        ],
    )
    add_para(doc, "Enum audit_action (9 valores):")
    add_bullet(doc, "login — signInWithPassword exitoso")
    add_bullet(doc, "login_failed — intento de auth fallido")
    add_bullet(doc, "logout — signOut")
    add_bullet(doc, "territory_toggle — apagar/prender desde admin")
    add_bullet(doc, "ptto_change — bulk upsert en territory_budgets")
    add_bullet(doc, "user_created — invite via Magic Link")
    add_bullet(doc, "user_updated — edit user (incluye reset password)")
    add_bullet(doc, "user_deleted — soft delete (is_active=false)")
    add_bullet(doc, "data_refresh — refresh APIs Susazón/Suve")

    add_h1(doc, "Vistas KPI")
    add_para(doc, "Todas las vistas con WITH (security_invoker = true) para heredar RLS de sales_rows.")
    add_table(
        doc,
        ["Vista", "Granularidad", "Tab que la consume"],
        [
            ["kpi_monthly_summary", "(anio, mes, territorio)", "KPI cards principales"],
            ["kpi_daily_summary", "(fecha, territorio)", "Tab Tracking Diario"],
            ["kpi_familia_summary", "(anio, mes, territorio, familia)", "Legacy"],
            ["kpi_grupo_summary", "(anio, mes, territorio, grupo)", "Tab Grupo Producto"],
            ["kpi_sku_summary", "(anio, mes, territorio, sku)", "Tab Productos"],
            ["kpi_cliente_summary", "(anio, mes, territorio, no_cliente)", "Tab Clientes"],
            ["kpi_vendedor_summary", "(anio, mes, territorio, vendedor, empresa)", "Tab Vendedores (toggle Sus/Suve)"],
            ["kpi_cliente_yearly", "(anio, no_cliente, territorio)", "Auxiliar"],
            ["kpi_cliente_perdidos", "(anio, no_cliente, territorio)", "Tab Perdidos (mes + YTD pre-agregados)"],
        ],
    )

    add_h1(doc, "Contrato APIs Susazón/Suve")
    add_h2(doc, "Endpoint")
    add_table(
        doc,
        ["Source", "URL", "Performance"],
        [
            ["Susazón", "https://sasweb.susazon.mx/susazon/api_ERPPyMEDashboard/", "SQL Enterprise · ~5s/mes"],
            ["Suve", "https://saswebsuve.susazon.mx/suve/api_ERPPyMEDashboard/", "SQL Express · ~60s/mes"],
        ],
    )

    add_h2(doc, "Request")
    add_code(doc, """POST /api_ERPPyMEDashboard/
Headers:
  X-API-KEY: API-DASH-CLAUDE-2026-$$1
  Content-Type: application/json

Body:
{
  "page": 1,
  "page_size": 50000,
  "desde": "2026-04-01",
  "hasta": "2026-04-30"
}""")

    add_h2(doc, "Response (filas)")
    add_table(
        doc,
        ["Campo", "Tipo", "Notas"],
        [
            ["empresa", "string", '"SUSAZON DEL CENTRO" | "SUVE DEL BAJIO" | etc — descriptivo, NO exacto'],
            ["no_cliente", "string", "ID del cliente"],
            ["cliente", "string", "Nombre del cliente"],
            ["territorio", "string", "Mapea 1-1 a sales_rows.territorio"],
            ["vendedor", "string", "Nombre del vendedor"],
            ["sku", "string", "Código de producto"],
            ["kg", "number | string", "Filtrar kg > 0 al insertar"],
            ["fecha", "string", "YYYY-MM-DD"],
            ["anio", "number | string", "Casteo a int"],
            ["mes", "number | string", "Casteo a int"],
            ["venta", "number | string", "MXN"],
            ["margen", "number | string", "MXN"],
            ["familia", "string | null", "Legacy"],
            ["grupo", "string | null", "Campo nuevo — no estaba en V2.2"],
        ],
    )
    add_para(doc, "Importante: campo `empresa` viene como STRING descriptivo. La normalización en lib/susazon-api.ts → normalizeRow() cae al fallback empresaCode (0 o 1 según el endpoint llamado).")

    add_h1(doc, "Estructura interna NormalizedRow (TypeScript)")
    add_code(
        doc,
        """interface NormalizedRow {
  empresa: 0 | 1;  // 0=Susazón, 1=Suve
  no_cliente: string;
  cliente: string | null;
  territorio: string;
  vendedor: string | null;
  sku: string | null;
  kg: number;
  fecha: string;
  anio: number;
  mes: number;
  venta: number;
  margen: number;
  familia: string | null;
  grupo: string | null;
}""",
    )

    add_h1(doc, "Roles y Matriz de Permisos")
    add_table(
        doc,
        ["Acción", "Admin", "Director", "Gerente Reg.", "Vendedor"],
        [
            ["Ver todos los territorios", "✅", "✅", "❌ (asignados)", "❌ (suyo)"],
            ["Editar PTTO", "✅", "✅", "❌", "❌"],
            ["Cargar datos (refresh APIs)", "✅", "✅", "❌", "❌"],
            ["Ver márgenes", "✅", "✅", "✅", "✅"],
            ["Apagar/prender territorios", "✅", "❌", "❌", "❌"],
            ["Crear/eliminar usuarios", "✅", "❌", "❌", "❌"],
            ["Ver audit log", "✅", "❌", "❌", "❌"],
        ],
    )

    add_h1(doc, "Constantes y Magic Numbers")
    add_table(
        doc,
        ["Constante", "Valor", "Razón"],
        [
            ["Susazón timeout/page", "120,000 ms (2 min)", "SQL Enterprise — margen amplio"],
            ["Suve timeout/page", "600,000 ms (10 min)", "SQL Express — lento por diseño"],
            ["PAGE_SIZE refresh", "50,000", "Bumpeado en Supabase"],
            ["maxDuration Vercel", "300s", "Hobby plan limit"],
            ["estimate Susazón", "5s/mes", "Observado en producción"],
            ["estimate Suve", "60s/mes", "Observado en producción"],
            ["acumulado años", "[2024, 2025, 2026]", "Configurable en page.tsx"],
            ["días hábiles", "L-S menos LFT", "No L-V"],
            ["min dia para Run-Rate", "5", "Antes hay poca data confiable"],
        ],
    )

    add_footer(doc, f"{PROYECTO} · Diccionario de Datos · {VERSION}")
    out = DOCS_DIR / "02_Diccionario_Datos.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")


# ============================================================
# DOC 03 — ChangeLog / Release Notes
# ============================================================
def gen_changelog():
    doc = Document()
    add_cover(doc, "ChangeLog & Release Notes", "Evolución V2.2 → V3.0 + historial de commits")

    add_h1(doc, "Versión 1.0.0 — Lanzamiento producción (2026-04-28)")
    add_para(doc, "Primera versión productiva del Dashboard V3.0. Reemplaza completamente al V2.2 con una arquitectura cliente-servidor segura, sistema de permisos y panel administrativo.", bold=False)

    add_h2(doc, "🚀 Nuevas features (vs V2.2)")
    add_bullet(doc, "Backend Next.js que oculta API keys del browser")
    add_bullet(doc, "Sistema de permisos por territorio + 4 roles (admin, director, gerente regional, vendedor)")
    add_bullet(doc, "Row-Level Security en Postgres como defensa en profundidad")
    add_bullet(doc, "Panel administrativo completo: gestión de usuarios, toggle de territorios, audit log")
    add_bullet(doc, "Audit log inmutable con 9 tipos de eventos capturados")
    add_bullet(doc, "Invitación de usuarios via Magic Link de Supabase (sin compartir passwords)")
    add_bullet(doc, "Editor de PTTO con persistencia en DB (no más localStorage del V2.2)")
    add_bullet(doc, "6 themes con selector modal (Clean, Editorial, Warm Neo, Susazón Moderno, Stock Market, Liquid Glass)")
    add_bullet(doc, "Theme Liquid Glass inspirado en iOS 26 Apple — aurora gradient + frosted glass")

    add_h2(doc, "🐛 Bugs del V2.2 corregidos")
    add_bullet(doc, "API Key Susazón ya NO es visible en el browser (vivía en HTML del V2.2)")
    add_bullet(doc, "Charts: eje X mostraba '$0, $1, $2...' en lugar de nombres → fix con XAxis dataKey='name' (Recharts)")
    add_bullet(doc, "Tab Familia → Grupo Producto: ahora grafica el campo `grupo` (TI lo agregó al API)")
    add_bullet(doc, "Tab Productos: agregado chart doble eje (Pesos + Kilos) arriba de la tabla")
    add_bullet(doc, "Tab Vendedores: separación correcta Sus/Suve con toggle (el V2.2 los mezclaba)")
    add_bullet(doc, "Tab Perdidos: toggle Mes/YTD con kg + monto en ambas dimensiones")

    add_h2(doc, "🔧 Mejoras técnicas")
    add_bullet(doc, "TypeScript estricto en todo el código")
    add_bullet(doc, "Migraciones SQL versionadas (10 migraciones aplicadas)")
    add_bullet(doc, "Refresh idempotente — seguro de re-correr cualquier rango")
    add_bullet(doc, "Vistas KPI pre-agregadas para performance")
    add_bullet(doc, "Manejo defensivo de env vars con sanitización de whitespace")

    add_h1(doc, "Historial de commits (32 totales)")
    add_para(doc, "Resumen de los commits más significativos en orden cronológico inverso:")

    add_h2(doc, "Cierre de proyecto + Deploy (2026-04-28)")
    add_table(
        doc,
        ["Hash", "Fecha", "Descripción"],
        [
            ["b81f182", "2026-04-28", "Fix Safari: backdrop-filter en Liquid Glass requiere GPU compositing"],
            ["f9a3eb6", "2026-04-28", "UX: warning visual cuando refresh excede límite Hobby de Vercel"],
            ["e00b51d", "2026-04-28", "Fix: validación defensiva de env vars + sanitize whitespace"],
            ["634c4ba", "2026-04-28", "Fix: maxDuration en /api/data/refresh dentro del límite Hobby (300s)"],
            ["3f9c1cf", "2026-04-28", "Fix: 3 errores de TypeScript que bloqueaban build de Vercel"],
        ],
    )

    add_h2(doc, "Themes Phase (2026-04-27 / 28)")
    add_table(
        doc,
        ["Hash", "Descripción"],
        [
            ["1c711c3", "Themes Chunk 4: upgrade ThemeSelector dropdown → modal con previews"],
            ["d46198a", "Themes Chunk 3: agregar theme 'Liquid Glass' (Apple iOS 26 + naranja)"],
            ["4d0700f", "Themes Chunk 2: agregar theme 'Stock Market' (trader desk)"],
            ["9d2eb30", "Themes Chunk 1: agregar theme 'Susazón Moderno' (dark + naranja)"],
        ],
    )

    add_h2(doc, "Admin Panel (Fase 5, 2026-04-28)")
    add_table(
        doc,
        ["Hash", "Descripción"],
        [
            ["c354fa6", "Fase 5 Chunk 3/3: Admin Panel - Audit Log read-only paginado"],
            ["6960c52", "Fase 5 Chunk 2/3: Admin Panel - Usuarios CRUD + magic link"],
            ["511caae", "Fase 5 Chunk 1/3: Admin Panel - Territorios toggle"],
        ],
    )

    add_h2(doc, "Dashboard 7 tabs (Fase 3, 2026-04-27)")
    add_table(
        doc,
        ["Hash", "Descripción"],
        [
            ["4b959a5", "Fase 3 Tab 6/6: Perdidos con toggle Mes / YTD (kg + monto)"],
            ["3e750dc", "Fase 3 Tab 5/6: Vendedores - agregar toggle Top 10 / Top 20 / Todos"],
            ["b7ae11b", "Fase 3 Tabs 4+5: Clientes + Vendedores (toggle Sus/Suve) + DimensionTab refactor"],
            ["fe25051", "Fase 3 Tab 3/6: Productos (cambio funcional #4 chart doble eje + toggle Top 10/15)"],
            ["067ee3f", "Fase 3 Tab 2/6: Grupo Producto + GroupedBarChart reusable + Migración 008"],
            ["457f8f4", "Fase 3 Tab 1/6: Ventas (replica V2.2 + custom tooltip)"],
        ],
    )

    add_h2(doc, "Fase 2 — Dashboard real")
    add_table(
        doc,
        ["Hash", "Descripción"],
        [
            ["bf0ea53", "AGENTS.md: snapshot al cierre Fase 2d"],
            ["9954230", "Fase 2d: Tab Tracking Diario completo (replica V2.2)"],
            ["fc53b2c", "Fase 2c Chunk 4: 6 KPI cards (Venta/Margen/KG con YoY + 3 Acum por anio)"],
            ["7746224", "Fase 2c Chunk 3: Editor de PTTO + fix Supabase 1000-row limit"],
        ],
    )

    add_h1(doc, "Migraciones de DB aplicadas")
    add_table(
        doc,
        ["Migración", "Propósito"],
        [
            ["001_users_permissions.sql", "Tabla users_permissions + helpers RLS"],
            ["002_territories_state.sql", "Tabla territories_state + visible_territories_for_current_user()"],
            ["003_audit_log.sql", "Tabla audit_log inmutable con enum de 9 acciones"],
            ["004_sales_rows.sql", "Tabla principal con RLS policy"],
            ["005_territory_budgets.sql", "PTTO mensual con compound PK"],
            ["006_kpi_views.sql", "kpi_monthly_summary y kpi_daily_summary"],
            ["007_kpi_daily_summary.sql", "Vista específica para Tracking Diario"],
            ["008_kpi_dimension_views.sql", "5 vistas de dimensión (familia, grupo, sku, vendedor, cliente)"],
            ["009_vendedor_with_empresa.sql", "Recreate kpi_vendedor_summary con empresa column"],
            ["010_kpi_cliente_perdidos.sql", "Vista para tab Perdidos (mes + YTD pre-agregados)"],
        ],
    )

    add_h1(doc, "Versiones futuras (planificación)")
    add_h2(doc, "v1.1.0 (propuesta)")
    add_bullet(doc, "Themes bonus: Linear Eclipse + Bento Spatial")
    add_bullet(doc, "Custom domain dashboard.susazon.mx")
    add_bullet(doc, "CI/CD con GitHub Actions (tests + lint)")

    add_h2(doc, "v1.2.0 (propuesta)")
    add_bullet(doc, "Sentry o LogRocket para monitoreo de errores runtime")
    add_bullet(doc, "Realtime updates en audit log (Supabase realtime)")
    add_bullet(doc, "Export PDF/Excel de los tabs")

    add_h2(doc, "v2.0.0 (futuro)")
    add_bullet(doc, "Reportes ejecutivos automáticos por email (cron)")
    add_bullet(doc, "Mobile responsive optimizado")
    add_bullet(doc, "Pricing dinámico por SKU")

    add_footer(doc, f"{PROYECTO} · ChangeLog · {VERSION}")
    out = DOCS_DIR / "03_ChangeLog_Release_Notes.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")


# ============================================================
# DOC 04 — Manual de Usuario
# ============================================================
def gen_manual():
    doc = Document()
    add_cover(doc, "Manual de Usuario", "Guía no-técnica para los 15 usuarios del sistema")

    add_h1(doc, "Bienvenida")
    add_para(doc, f"Bienvenido al Dashboard Comercial V3.0 de Grupo Susazón. Esta es una herramienta interna que te permite ver datos comerciales en tiempo real (ventas, márgenes, kilos, clientes, productos, vendedores) según los territorios que tu rol te permite consultar.")
    add_para(doc, "El acceso es exclusivo para 15 personas autorizadas. Toda la actividad queda registrada en el sistema de auditoría.")

    add_h1(doc, "Primeros pasos")
    add_h2(doc, "1. Cómo entrar")
    add_para(doc, "Vas a recibir un email de invitación de Supabase con un Magic Link. Hacé click en el link para configurar tu contraseña la primera vez. Después podés entrar normal con email + contraseña en la URL del dashboard.")
    add_bullet(doc, "URL: la que te pase Mauricio (algo terminado en .vercel.app)")
    add_bullet(doc, "Si olvidás la contraseña: click en '¿Olvidaste tu contraseña?' en la pantalla de login")
    add_bullet(doc, "Si tu email rebota o no llega: contactá a Mauricio para reenvío")

    add_h2(doc, "2. La pantalla principal")
    add_para(doc, "Cuando entrás vas a ver:")
    add_bullet(doc, "Header arriba: logo Susazón, navegación, tu nombre y rol, selector de tema, botón de cerrar sesión")
    add_bullet(doc, "Sidebar izquierdo: lista de territorios que podés ver. Click sobre uno para filtrar todo el dashboard a ese territorio")
    add_bullet(doc, "Centro: 6 cards de KPIs (Venta, Margen, KG, Acum 2024/2025/2026)")
    add_bullet(doc, "Tabs abajo: 7 secciones distintas con análisis específicos")

    add_h2(doc, "3. Cambiar de tema")
    add_para(doc, "Click en el botón 'Tema' (paleta) en el header. Se abre un modal con 6 opciones — cada una con un preview real del look. Click en la que más te late y se aplica al instante. Tu elección queda guardada en tu navegador.")
    add_para(doc, "Themes disponibles:", bold=True)
    add_bullet(doc, "Dashboard Clean: corporativo, blanco/azul, similar al V2.2")
    add_bullet(doc, "Susazón Editorial: elegante, fuente Bebas Neue, naranja sobre paper warm")
    add_bullet(doc, "Warm Neo-Editorial: minimalista crema + coral, estilo Anthropic")
    add_bullet(doc, "Susazón Moderno: dark back-office con naranja Susazón, ideal para análisis denso")
    add_bullet(doc, "Stock Market: trader desk azul nocturno + neón cyan/verde, estilo Bloomberg")
    add_bullet(doc, "Liquid Glass: Apple iOS 26 con aurora naranja + frosted glass — moderno y atractivo")

    add_h1(doc, "Los 7 tabs explicados")

    add_h2(doc, "Tab 1: Tracking Diario")
    add_para(doc, "Muestra el avance del mes actual día por día.")
    add_bullet(doc, "8 stats arriba: Venta del Mes, Alcance PTTO, Margen, vs Mismo Mes Año Anterior, Vel. Original/Actual/Necesaria, Run Rate")
    add_bullet(doc, "Progress bar verde/amarilla/roja según vayas vs el PTTO y el tiempo del mes")
    add_bullet(doc, "Chart con 4 series: Acumulado actual, Ptto Linear, Año Anterior, Venta Diaria")
    add_bullet(doc, "Tabla diaria con 9 columnas: día, día semana, venta, margen, kg, acum, brecha, vel necesaria")
    add_bullet(doc, "Color de Vel. Necesaria: verde si vas bien, amarillo intermedio, rojo si te exigirá esfuerzo")

    add_h2(doc, "Tab 2: Ventas")
    add_para(doc, "Comparativo histórico de los últimos 12 meses × 3 años (2024, 2025, 2026).")
    add_bullet(doc, "Bars: venta en pesos por mes")
    add_bullet(doc, "Lines: margen % por mes")
    add_bullet(doc, "Tooltip rico al pasar el cursor: muestra YoY delta y porcentajes")

    add_h2(doc, "Tab 3: Grupo Producto")
    add_para(doc, "Análisis por grupo de producto (campo nuevo agregado por TI).")
    add_bullet(doc, "Bar chart: Top 10 grupos por venta del mes actual, comparados contra mismo mes 2025 y 2024")
    add_bullet(doc, "Tabla: todos los grupos con venta de los 3 años + Var %")

    add_h2(doc, "Tab 4: Productos (SKUs)")
    add_para(doc, "Análisis al detalle de SKUs.")
    add_bullet(doc, "Chart doble eje arriba: barras de Kilos (eje izq) + líneas de Pesos (eje der)")
    add_bullet(doc, "Toggle 'Top 10' / 'Top 15' para ajustar el detalle del chart")
    add_bullet(doc, "Tabla abajo: Top 50 SKUs con venta y kilos en los 3 años")

    add_h2(doc, "Tab 5: Clientes")
    add_para(doc, "Top 50 clientes por venta del mes actual.")
    add_bullet(doc, "Bar chart: Top 10 clientes, comparados 24/25/26")
    add_bullet(doc, "Tabla con detalle y % de variación")

    add_h2(doc, "Tab 6: Vendedores")
    add_para(doc, "Performance individual de la fuerza de ventas.")
    add_bullet(doc, "Toggle 'Separar Sus / Suve' o 'Unir Sus + Suve' — decide si los vendedores que venden ambas marcas se muestran como filas distintas o agregados")
    add_bullet(doc, "Toggle Top 10 / Top 20 / Todos")
    add_bullet(doc, "Bar chart + tabla con Var %")

    add_h2(doc, "Tab 7: Perdidos")
    add_para(doc, "Clientes perdidos o en declive.")
    add_bullet(doc, "Toggle 'Mes Actual' / 'YTD' para ver qué clientes perdimos en el mes vs en lo que va del año")
    add_bullet(doc, "3 stats: Total Perdidos, Declive >30%, Total Declive")
    add_bullet(doc, "Tabla 9 columnas: Cliente, Vendedor, Status (Perdido/Declive), $25, $26, Var $%, kg25, kg26, Var kg%")
    add_bullet(doc, "Color rojo: declive >30% o total. Amarillo: declive <30%. Caso especial: kg subió pero $ bajó (vendió más cantidad pero menos dinero — bajó precios)")

    add_h1(doc, "Si sos Admin")
    add_para(doc, "Verás un botón 'Admin' (escudo) en el header. Click te lleva al panel administrativo con 3 sub-secciones:")

    add_h2(doc, "Territorios")
    add_para(doc, "Tabla con los 17 territorios. Podés apagar un territorio (ej. cuando no es comercial activo) — al hacerlo, NADIE en la app va a verlo, sin importar sus permisos. Apagar pide motivo opcional.")
    add_para(doc, "Cada toggle queda registrado en audit log con tu nombre y timestamp.")

    add_h2(doc, "Usuarios")
    add_para(doc, "Tabla con todos los 15 usuarios. Acciones disponibles:")
    add_bullet(doc, "Botón 'Invitar usuario' — abre modal con email, nombre, rol, territorios permitidos. Supabase manda Magic Link automático")
    add_bullet(doc, "Editar (lápiz) — cambiar rol, territorios, can_edit_ptto")
    add_bullet(doc, "Reset password (llave) — manda email de reset al user")
    add_bullet(doc, "Apagar/encender (power) — desactivar sin borrar (soft delete)")
    add_para(doc, "Por seguridad, no podés desactivarte a vos mismo ni bajarte de rol admin.")

    add_h2(doc, "Audit Log")
    add_para(doc, "Tabla paginada con todos los eventos de seguridad y cambios sensibles. Filtros por:")
    add_bullet(doc, "Acción (login, logout, territory_toggle, ptto_change, user_*, data_refresh)")
    add_bullet(doc, "Email (búsqueda parcial)")
    add_bullet(doc, "Rango de fechas (desde / hasta)")
    add_para(doc, "Click en el chevron ▶ de cada fila para expandir el JSON completo del evento (incluye snapshot before/after en updates).")

    add_h1(doc, "Si sos Admin o Director")
    add_para(doc, "Verás también el botón 'Cargar datos' en el header. Va a /cargar-datos donde podés:")

    add_h2(doc, "Refrescar datos desde APIs")
    add_para(doc, "Selecciona rango de meses (Desde / Hasta) y fuentes (Susazón / Suve). Click 'Iniciar refresh'.")
    add_para(doc, "⚠️ Importante: el sistema te muestra un estimado de tiempo. Si pedís un rango muy largo con ambas fuentes, vas a ver un warning amarillo. Razón: el plan actual de Vercel limita las funciones a 5 minutos. Si pedís más de 5 meses con Suve, va a fallar. Refrescá rangos chicos (1-3 meses) o desmarcá Suve para rangos largos.")
    add_para(doc, "Tu data histórica YA está en la base — solo necesitás refrescar el mes actual o los últimos 1-2 meses para tener los datos frescos.")

    add_h2(doc, "Editar PTTO (Presupuesto)")
    add_para(doc, "Grid editable año × territorio × 12 meses. Selecciona el año y editá las celdas. Hay autoguardado — los cambios se persisten automáticamente.")

    add_h1(doc, "FAQ")
    add_h2(doc, "¿Por qué no veo todos los territorios?")
    add_para(doc, "Solo ves los territorios que tu rol y tu lista de allowed_territories permite. Si esperás ver más, contactá a Mauricio para ajustar tus permisos.")

    add_h2(doc, "¿Por qué un territorio aparece sombreado en el sidebar?")
    add_para(doc, "Está temporalmente apagado por el administrador (vas a ver un banner amarillo en la parte superior con el motivo). No vas a poder consultar sus datos hasta que se prenda de nuevo.")

    add_h2(doc, "¿Qué hago si un dato no cuadra con el ERP?")
    add_para(doc, "Verificá la fecha y la última sincronización. Si ya pasó tiempo desde el último refresh, decile a Mauricio o a un director que dispare un refresh nuevo.")

    add_h2(doc, "¿Cómo cierro mi sesión?")
    add_para(doc, "Click en 'Cerrar sesión' en el header arriba a la derecha. Tu sesión se invalida y vas al login.")

    add_h2(doc, "¿Funciona en mobile?")
    add_para(doc, "Funciona en pantallas grandes. En mobile algunos charts y tablas pueden verse apretados — está optimizado para desktop y tablets en horizontal.")

    add_h2(doc, "¿Qué browser uso?")
    add_para(doc, "Chrome, Safari, Edge o Firefox actualizados. El theme Liquid Glass requiere browser moderno (2023+) para verse bien con el frosted glass.")

    add_footer(doc, f"{PROYECTO} · Manual de Usuario · {VERSION}")
    out = DOCS_DIR / "04_Manual_Usuario.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")


# ============================================================
# DOC 05 — Guía TI / Despliegue
# ============================================================
def gen_guia_ti():
    doc = Document()
    add_cover(doc, "Guía de TI y Despliegue", "Para ingenieros que mantienen, deployan o continúan el sistema")

    add_h1(doc, "Quién debe leer esto")
    add_para(doc, "Este documento está pensado para un ingeniero de TI que necesite:")
    add_bullet(doc, "Setear el proyecto desde cero en una nueva máquina")
    add_bullet(doc, "Hacer cambios en el código y deployarlos a producción")
    add_bullet(doc, "Diagnosticar bugs en runtime")
    add_bullet(doc, "Migrar el proyecto a otra cuenta (Vercel, Supabase, GitHub)")
    add_bullet(doc, "Continuar el desarrollo de features futuras")

    add_h1(doc, "Prerequisitos")
    add_h2(doc, "Software local")
    add_table(
        doc,
        ["Herramienta", "Versión mínima", "Cómo instalar (macOS)"],
        [
            ["Node.js", "20.x", "brew install node@20 (o nvm)"],
            ["npm", "10.x", "incluido con Node"],
            ["Git", "2.30+", "incluido en macOS"],
            ["GitHub CLI (opcional)", "2.x", "brew install gh"],
            ["Editor", "VS Code recomendado", "code.visualstudio.com"],
        ],
    )

    add_h2(doc, "Cuentas y accesos necesarios")
    add_bullet(doc, "GitHub: acceso al repo dashboard-susazon-v3 (privado, solicitar a Mauricio)")
    add_bullet(doc, "Vercel: invite de Mauricio al proyecto Vercel")
    add_bullet(doc, "Supabase: invite de Mauricio al proyecto qfxyrpifntcixwpvnjpd")
    add_bullet(doc, "Acceso al archivo SECRETS_DASHBOARD_V3.txt (Apple Notes de Mauricio)")

    add_h1(doc, "Setup local desde cero")
    add_h2(doc, "1. Clonar el repo")
    add_code(doc, """git clone https://github.com/musabiaga/dashboard-susazon-v3.git
cd dashboard-susazon-v3""")

    add_h2(doc, "2. Configurar variables de entorno")
    add_para(doc, "Copiar .env.example a .env.local y completar los valores con los del archivo SECRETS:")
    add_code(doc, """cp .env.example .env.local
# Editar .env.local y pegar los valores del SECRETS_DASHBOARD_V3.txt:
# - NEXT_PUBLIC_SUPABASE_URL
# - NEXT_PUBLIC_SUPABASE_ANON_KEY
# - SUPABASE_SERVICE_ROLE_KEY
# - SUSAZON_API_URL
# - SUSAZON_API_KEY (escapar $$ como \\$\\$)
# - SUVE_API_URL
# - SUVE_API_KEY (escapar $$ como \\$\\$)""")
    add_para(doc, "⚠️ IMPORTANTE: en .env.local los $$ se escriben como \\$\\$ porque dotenv-expand los interpreta como variables.", bold=True)
    add_para(doc, "En Vercel UI, en cambio, los valores se pegan literales sin escape.")

    add_h2(doc, "3. Instalar dependencias")
    add_code(doc, "npm install")

    add_h2(doc, "4. Verificar el build")
    add_code(doc, """# TypeScript check
npx tsc --noEmit

# Build de producción (debe compilar 19 rutas)
npm run build

# Si pasa, dev local
npm run dev""")
    add_para(doc, "El dashboard debe estar accesible en http://localhost:3000")

    add_h1(doc, "Estructura de variables de entorno")
    add_table(
        doc,
        ["Variable", "Visibilidad", "Propósito"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL", "Pública (browser OK)", "URL del proyecto Supabase"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Pública", "Anon key de Supabase. Permisos limitados por RLS"],
            ["SUPABASE_SERVICE_ROLE_KEY", "PRIVADA", "Bypassea RLS. Solo en API routes. JAMÁS al cliente"],
            ["SUSAZON_API_URL", "PRIVADA", "Endpoint REST Susazón"],
            ["SUSAZON_API_KEY", "PRIVADA", "X-API-KEY header (escapar $$)"],
            ["SUVE_API_URL", "PRIVADA", "Endpoint REST Suve"],
            ["SUVE_API_KEY", "PRIVADA", "Misma key que Susazón"],
        ],
    )

    add_h1(doc, "Deploy a producción")
    add_h2(doc, "Setup inicial (solo primera vez)")
    add_bullet(doc, "Crear cuenta Vercel con el GitHub que tiene acceso al repo")
    add_bullet(doc, "Add New Project → Import dashboard-susazon-v3")
    add_bullet(doc, "Framework Preset: Next.js (auto-detectado)")
    add_bullet(doc, "Pegar las 7 env vars en la sección Environment Variables")
    add_bullet(doc, "Click Deploy")

    add_h2(doc, "Deploy de cambios subsecuentes")
    add_para(doc, "Vercel deploya automáticamente en cada push a main. Workflow:")
    add_code(doc, """# 1. Hacer cambios
# 2. Verificar tipos
npx tsc --noEmit

# 3. Verificar build local (CRÍTICO — Vercel corre tsc strict)
npm run build

# 4. Commit
git add .
git commit -m "Descripción clara del cambio

Detalles del por qué.

Co-Authored-By: ..."

# 5. Push (Vercel detecta y deploya en ~3 min)
git push origin main""")

    add_h1(doc, "Variables de entorno en Vercel")
    add_para(doc, "Settings → Environment Variables. Confirmar que las 7 están seteadas para Production / Preview / Development. Si alguna tiene whitespace o caracteres invisibles al pegar, va a fallar el refresh con 'string did not match expected pattern'.")

    add_h1(doc, "Limitaciones del plan Hobby")
    add_table(
        doc,
        ["Recurso", "Hobby", "Pro ($20/mes)"],
        [
            ["maxDuration por función", "300s (5 min)", "900s (15 min)"],
            ["Build minutes/mes", "6,000", "24,000"],
            ["Bandwidth/mes", "100 GB", "1 TB"],
            ["Domains", "1 .vercel.app", "Unlimited custom"],
            ["Team members", "1", "Unlimited"],
        ],
    )
    add_para(doc, "Si Mauricio quiere refrescar más de 5 meses con ambas APIs (>5 min de procesamiento), o quiere custom domain dashboard.susazon.mx con SSL automático, considerar upgrade a Pro.")

    add_h1(doc, "Migraciones de base de datos")
    add_para(doc, "Las 10 migraciones SQL viven en supabase/migrations/. Si cambia el schema:")
    add_h2(doc, "1. Crear archivo numerado")
    add_code(doc, "supabase/migrations/011_descripcion_cambio.sql")

    add_h2(doc, "2. Aplicarla")
    add_para(doc, "Hay 2 formas:")
    add_bullet(doc, "Manual: copiar el SQL al SQL Editor en supabase.com → ejecutar")
    add_bullet(doc, "CLI (preferido si tenés supabase CLI instalado): supabase db push")

    add_h2(doc, "3. Documentar")
    add_para(doc, "Actualizar SESSION_LOG.md con la decisión y el motivo. Actualizar 02_Diccionario_Datos.docx en el próximo regen de docs.")

    add_h1(doc, "Troubleshooting")
    add_h2(doc, "Build de Vercel falla con TypeScript")
    add_para(doc, "Síntoma: 'Type X is not assignable to type Y'.")
    add_para(doc, "Causa común: tipos de Recharts cambiaron. npm run dev no corre tsc strict, Vercel sí.")
    add_para(doc, "Fix: correr npx tsc --noEmit local, fixear cada error, recommit.")

    add_h2(doc, "API refresh hace timeout")
    add_para(doc, "Síntoma: 'Vercel Runtime Timeout Error: Task timed out after 300 seconds'.")
    add_para(doc, "Causa: rango muy largo + Suve marcado.")
    add_para(doc, "Fix: rangos de 1-3 meses o desmarcar Suve. O upgrade a Pro plan.")

    add_h2(doc, "Refresh dice 'string did not match expected pattern'")
    add_para(doc, "Síntoma: error críptico al darle Iniciar refresh.")
    add_para(doc, "Causa: env var con whitespace al copy/paste en Vercel UI.")
    add_para(doc, "Fix: ir a Settings → Environment Variables → click ojo 👁️ en SUSAZON_API_URL → confirmar formato exacto sin espacios. lib/susazon-api.ts ahora sanitiza, pero igual conviene corregir el origen.")

    add_h2(doc, "Sidebar muestra solo 1 territorio")
    add_para(doc, "Causa: query a sales_rows hit el límite de 1000 filas de Supabase.")
    add_para(doc, "Fix: ya bumpeado a 50000. Si vuelve, ir a Supabase → Settings → API → Max rows.")

    add_h2(doc, "Theme Liquid Glass se ve raro en Safari")
    add_para(doc, "Síntoma: cards sólidas en vez de blureadas.")
    add_para(doc, "Causa: Safari requiere isolation: isolate + transform: translateZ(0) para backdrop-filter.")
    add_para(doc, "Fix: ya aplicado en globals.css. Si vuelve a fallar, verificar que selector [class*='rounded-'] siga matcheando los elementos correctos.")

    add_h2(doc, "Turbopack cachea CSS")
    add_para(doc, "Síntoma: cambios en globals.css no aparecen en dev.")
    add_para(doc, "Fix: rm -rf .next && npm run dev")

    add_h1(doc, "Monitoreo")
    add_h2(doc, "Vercel Logs")
    add_para(doc, "Vercel Dashboard → tu proyecto → Logs. Filtrá por path (/api/data/refresh, /api/admin/*) para ver errores.")

    add_h2(doc, "Supabase Logs")
    add_para(doc, "Supabase Dashboard → Logs → Database / Auth / API. Util para diagnosticar RLS issues.")

    add_h2(doc, "Audit log de la app")
    add_para(doc, "Mauricio (admin) puede ir a /admin/audit en el dashboard mismo para ver eventos de seguridad: login attempts, cambios de PTTO, toggles de territorio, gestión de usuarios.")

    add_h1(doc, "Backup y recovery")
    add_h2(doc, "Supabase backups")
    add_para(doc, "Free tier: backup automático cada 24h, retention 7 días. Pro: 14 días + Point-in-Time Recovery.")
    add_para(doc, "Para snapshot manual: Supabase Dashboard → Database → Backups → 'Create backup'.")

    add_h2(doc, "Repositorio")
    add_para(doc, "GitHub guarda el código. Para mirror local periódico:")
    add_code(doc, "git clone --mirror https://github.com/musabiaga/dashboard-susazon-v3.git ~/Backups/dashboard-susazon-v3-mirror.git")

    add_h1(doc, "Seguridad operacional")
    add_h2(doc, "Rotación de tokens")
    add_para(doc, "Rotar la SUPABASE_SERVICE_ROLE_KEY cada 6 meses (ir a Supabase → Settings → API → Reset). Después updatear en Vercel env vars.")
    add_para(doc, "Las API keys de Susazón/Suve las controla TI de Susazón — coordinar con ellos.")

    add_h2(doc, "Permisos de GitHub")
    add_para(doc, "El repo es privado. Solo Mauricio (owner) puede dar acceso. Para colaboradores temporales (consultores, devs externos), preferir invites por tiempo definido.")

    add_h2(doc, "Personal Access Tokens")
    add_para(doc, "Si necesitás push desde una máquina sin SSH, usar Fine-grained PAT con scope mínimo (Contents: Read+Write, solo para este repo). Revocar al terminar la tarea.")

    add_footer(doc, f"{PROYECTO} · Guía TI · {VERSION}")
    out = DOCS_DIR / "05_Guia_TI_Despliegue.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")


# ============================================================
# DOC 06 — Guía de Reconstrucción
# ============================================================
def gen_reconstruccion():
    doc = Document()
    add_cover(doc, "Guía de Reconstrucción", "Cómo rebuildear el sistema desde cero")

    add_h1(doc, "Cuándo usar esta guía")
    add_para(doc, "Esta es la guía CRÍTICA — debe permitir a un ingeniero competente reconstruir el sistema completo desde cero, asumiendo que solo tiene esta documentación + acceso a las APIs externas.")
    add_para(doc, "Casos de uso:")
    add_bullet(doc, "Migrar a otra cuenta de Supabase / Vercel / GitHub")
    add_bullet(doc, "Recuperar de un incidente catastrófico")
    add_bullet(doc, "Replicar el sistema para un cliente nuevo (con adaptaciones)")
    add_bullet(doc, "Auditar la implementación contra una especificación")

    add_h1(doc, "Etapa 1 — Setup de infraestructura externa")
    add_h2(doc, "1.1 Crear proyecto Supabase")
    add_bullet(doc, "Ir a supabase.com → New Project")
    add_bullet(doc, "Region: East US (N. Virginia) — match con la actual")
    add_bullet(doc, "Plan: Free tier es suficiente. Bumpear max-rows a 50000 en Settings → API")
    add_bullet(doc, "Copiar URL, anon key, service_role key del Settings → API")

    add_h2(doc, "1.2 Aplicar las 10 migraciones SQL")
    add_para(doc, "Las migraciones viven en supabase/migrations/. Aplicarlas en orden 001 → 010 desde el SQL Editor de Supabase. Cada una debe ejecutarse limpia sin errores.")
    add_para(doc, "Verificación post-migración:")
    add_code(doc, """-- Debe retornar las 5 tablas core
SELECT table_name FROM information_schema.tables
WHERE table_schema = 'public'
ORDER BY table_name;
-- Esperado: audit_log, sales_rows, territories_state,
-- territory_budgets, users_permissions

-- Debe retornar las 9 vistas con security_invoker = true
SELECT relname,
  (SELECT option_value FROM pg_options_to_table(reloptions)
   WHERE option_name = 'security_invoker') AS security_invoker
FROM pg_class
WHERE relkind = 'v' AND relname LIKE 'kpi_%';""")

    add_h2(doc, "1.3 Crear cuenta Vercel + repo GitHub")
    add_bullet(doc, "GitHub: crear repo privado con el contenido del backup")
    add_bullet(doc, "Vercel: crear cuenta con GitHub linked")
    add_bullet(doc, "Vercel: import del repo, framework Next.js auto-detectado")

    add_h1(doc, "Etapa 2 — Preparar variables de entorno")
    add_para(doc, "Recolectar 7 valores:")
    add_table(
        doc,
        ["Variable", "Origen"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL", "Supabase → Settings → API → Project URL"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Supabase → Settings → API → anon public"],
            ["SUPABASE_SERVICE_ROLE_KEY", "Supabase → Settings → API → service_role secret"],
            ["SUSAZON_API_URL", "Susazón TI: https://sasweb.susazon.mx/.../"],
            ["SUSAZON_API_KEY", "Susazón TI provee. Format API-DASH-CLAUDE-YYYY-XXX"],
            ["SUVE_API_URL", "Susazón TI: https://saswebsuve.susazon.mx/.../"],
            ["SUVE_API_KEY", "Susazón TI: misma que Susazón"],
        ],
    )
    add_para(doc, "Pegar en Vercel UI (Production + Preview + Development) y en .env.local local (con escape \\$\\$ en las API keys).")

    add_h1(doc, "Etapa 3 — Crear admin inicial")
    add_para(doc, "Sin un admin no podés gestionar nada. Pasos:")
    add_h2(doc, "3.1 Crear cuenta Supabase Auth")
    add_bullet(doc, "Supabase → Authentication → Users → Invite user con email del admin")
    add_bullet(doc, "Aceptar invite + setear password")
    add_bullet(doc, "Anotar el UUID generado (Authentication → Users → click row → ID)")

    add_h2(doc, "3.2 Insertar permisos admin")
    add_code(doc, """INSERT INTO public.users_permissions
  (user_id, email, full_name, role, allowed_territories, can_edit_ptto, is_active)
VALUES (
  '<UUID_DE_PASO_3.1>',
  'admin@empresa.com.mx',
  'Nombre del Admin',
  'admin',
  NULL,  -- NULL = ve todos los territorios
  true,
  true
);""")

    add_h1(doc, "Etapa 4 — Cargar datos iniciales")
    add_h2(doc, "4.1 Deploy primer build a Vercel")
    add_bullet(doc, "Push del código a GitHub main branch")
    add_bullet(doc, "Vercel detecta y hace primer deploy (~3 min)")
    add_bullet(doc, "Si falla: revisar logs, fixear, re-push")

    add_h2(doc, "4.2 Login y refresh inicial")
    add_bullet(doc, "Loguear como admin en la URL .vercel.app/login")
    add_bullet(doc, "Ir a /cargar-datos")
    add_bullet(doc, "Refresh inicial: rango chico (último mes) con ambas fuentes para validar conectividad")
    add_bullet(doc, "Si funciona, ir ampliando: 3 meses, 6 meses, etc., respetando el límite de 5 min del Hobby plan")

    add_h2(doc, "4.3 Cargar PTTO histórico")
    add_para(doc, "Si tenés un sheet con PTTO de meses pasados:")
    add_bullet(doc, "Ir a /cargar-datos → sección Editor PTTO")
    add_bullet(doc, "Seleccionar año, ingresar valores en el grid territorio × mes")
    add_bullet(doc, "Autoguardado: cada celda se persiste cuando perdés focus")
    add_bullet(doc, "Alternativamente: bulk INSERT vía SQL si tenés muchas filas (más rápido)")

    add_h1(doc, "Etapa 5 — Invitar usuarios")
    add_para(doc, "Una vez con datos, invitar al resto del equipo:")
    add_bullet(doc, "Ir a /admin/usuarios → 'Invitar usuario'")
    add_bullet(doc, "Email + nombre + rol + allowed_territories + can_edit_ptto")
    add_bullet(doc, "Supabase manda Magic Link al usuario")
    add_bullet(doc, "Usuario configura su password al recibir el email")
    add_para(doc, "Importante: Supabase Free tier tiene rate limit de ~3 emails/hora. Para invitar a los 15 usuarios de una vez, espaciar los invites o configurar SMTP custom (Resend/SendGrid).")

    add_h1(doc, "Algoritmos críticos")

    add_h2(doc, "Tracking Diario — fórmulas")
    add_para(doc, "Estas fórmulas son la lógica de negocio del Tab Tracking. Portadas verbatim del V2.2.")
    add_code(doc, """// Días hábiles del mes (Lunes a Sábado, excluyendo feriados LFT)
totalBizDays = countBizDays(year, month, null);
elapsedBizDays = countBizDays(year, month, dayOfMonth);
remainingBizDays = totalBizDays - elapsedBizDays;

// Velocidades
velOrig = ptto / totalBizDays;            // velocidad original requerida
velActual = acum / elapsedBizDays;        // velocidad actual lograda
velNeces = (ptto - acum) / remainingBizDays; // velocidad para alcanzar el ptto
runRate = velActual * totalBizDays;       // proyección lineal

// Indicadores
alcancePct = (acum / ptto) * 100;
tiempoPct = (elapsedBizDays / totalBizDays) * 100;
brecha = alcancePct - tiempoPct;
estado = brecha >= 0 ? "AVANZADO" : "REZAGADO";

// Color tones para Vel.Necesaria
if (velNeces <= velOrig) color = "green";
else if (velNeces <= velOrig * 1.2) color = "yellow";
else color = "red";""")

    add_h2(doc, "Días hábiles — tabla de feriados LFT")
    add_para(doc, "lib/business-days.ts tiene la tabla hardcoded de feriados Ley Federal del Trabajo 2024-2027. Es L-S menos Domingos menos feriados. Actualizar cada año.")
    add_para(doc, "Ejemplos 2026:")
    add_bullet(doc, "1 Ene (Año Nuevo)")
    add_bullet(doc, "2 Feb 1er Lun (Constitución)")
    add_bullet(doc, "16 Mar 3er Lun (Natalicio Juárez)")
    add_bullet(doc, "1 May (Trabajo)")
    add_bullet(doc, "16 Sep (Independencia)")
    add_bullet(doc, "16 Nov 3er Lun (Revolución)")
    add_bullet(doc, "25 Dic (Navidad)")

    add_h2(doc, "Refresh idempotente")
    add_para(doc, "Para cada (source, año, mes):")
    add_code(doc, """1. fetchMonth(source, year, month) → array de NormalizedRow (kg > 0)
2. supabase
     .from('sales_rows')
     .delete()
     .match({ empresa: source==='susazon'?0:1, anio: year, mes: month });
3. supabase
     .from('sales_rows')
     .insert(rows);
4. audit_log entry con sync_id, rows_imported, etc.""")
    add_para(doc, "El DELETE antes del INSERT garantiza que re-correr el mismo rango no produzca duplicados.")

    add_h2(doc, "RLS — Visibilidad de territorios")
    add_para(doc, "La función helper que decide qué territorios ve un usuario:")
    add_code(doc, """CREATE FUNCTION visible_territories_for_current_user()
RETURNS text[] AS $$
DECLARE
  user_terr text[];
  active_terr text[];
BEGIN
  -- Allowed_territories del user (NULL = todos los activos)
  SELECT allowed_territories INTO user_terr
  FROM users_permissions WHERE user_id = auth.uid();

  -- Territorios actualmente activos
  SELECT array_agg(territory_name) INTO active_terr
  FROM territories_state WHERE is_active = true;

  IF user_terr IS NULL THEN
    RETURN COALESCE(active_terr, ARRAY[]::text[]);
  END IF;

  -- Intersección
  RETURN ARRAY(
    SELECT unnest(user_terr)
    INTERSECT
    SELECT unnest(COALESCE(active_terr, ARRAY[]::text[]))
  );
END;
$$ LANGUAGE plpgsql STABLE SECURITY DEFINER;""")

    add_h2(doc, "Tab Perdidos — clasificación")
    add_code(doc, """if (v25 > 0 && v26 === 0) {
  status = "perdido";
  declinePct = 100;
} else if (v25 > 0 && v26 < v25) {
  status = "declive";
  declinePct = ((v25 - v26) / v25) * 100;
} else {
  status = null; // No es perdido ni declive
}""")

    add_h1(doc, "Verificación post-rebuild")
    add_para(doc, "Checklist para confirmar que el rebuild es funcional:")

    add_h2(doc, "Seguridad")
    add_bullet(doc, "Inspeccionar bundle JS desde DevTools — confirmar que SUSAZON_API_KEY NO aparece")
    add_bullet(doc, "Confirmar que SUPABASE_SERVICE_ROLE_KEY NO aparece en el bundle")
    add_bullet(doc, "Crear user 'vendedor' con 1 solo territorio → loguear → confirmar que solo ve ese territorio")
    add_bullet(doc, "Como vendedor, hacer fetch directo a /api/data/snapshot de OTRO territorio → debe retornar 403")
    add_bullet(doc, "Apagar un territorio desde admin → confirmar banner amarillo + sidebar shading")

    add_h2(doc, "Funcionalidad")
    add_bullet(doc, "Login funciona con admin")
    add_bullet(doc, "Dashboard carga con KPIs y los 7 tabs")
    add_bullet(doc, "Refresh APIs trae filas (verificar conteo en Supabase)")
    add_bullet(doc, "Editor PTTO autosaves (verificar persistencia recargando)")
    add_bullet(doc, "Admin Panel: invitar usuario → Magic Link llega → usuario puede entrar")
    add_bullet(doc, "Admin Panel: toggle territorio aparece en audit log")
    add_bullet(doc, "Audit log paginado funciona con filtros")
    add_bullet(doc, "Los 6 themes cambian apariencia de toda la app")
    add_bullet(doc, "Theme Liquid Glass se ve correcto en Safari (no solo Chrome)")

    add_h2(doc, "Datos")
    add_bullet(doc, "Comparar venta total enero 2025 vs el sistema viejo (V2.2 si existe) → diferencia <0.1%")
    add_bullet(doc, "Sumar manual los KPIs de los 16 territorios + los apagados = total mostrado en 'Todos'")

    add_h1(doc, "Gotchas conocidos")
    add_para(doc, "Lista de detalles que NO se descubren leyendo solo el código — vienen de horas de debugging:")

    add_table(
        doc,
        ["Gotcha", "Solución"],
        [
            ["middleware.ts no funciona en Next 16", "Renombrar a proxy.ts con export function proxy()"],
            ["cookies() es async en Next 16", "await cookies()"],
            ["dotenv-expand interpola $$", "Escapar como \\$\\$ en .env.local"],
            ["empresa viene STRING del API, no int", "normalizeRow() mapea string→0|1 con fallback empresaCode"],
            ["Supabase 1000-row limit", "Bumpear a 50000 en Settings → API"],
            ["Vercel Hobby 300s limit", "maxDuration <= 300, warning en UI"],
            ["Recharts 3.x payload readonly", "Cambiar interface a readonly TooltipItem[]"],
            ["Safari backdrop-filter no aplica", "isolation: isolate + transform: translateZ(0)"],
            ["Turbopack cachea CSS", "rm -rf .next entre cambios de globals.css"],
            ["Carpetas con _ no se rutean en App Router", "Usar - o nombre normal (_debug → debug)"],
            ["Vercel SSO bloquea acceso público", "Settings → Deployment Protection → Off"],
            ["Suve API 60s/mes timeout", "timeoutMs: 600_000 en getApiConfig"],
        ],
    )

    add_h1(doc, "Dependencias con versiones exactas")
    add_para(doc, "Para reproducibilidad bit-exacta, fijar:")
    add_table(
        doc,
        ["Package", "Versión exacta"],
        [
            ["next", "16.2.4"],
            ["react", "19.2.4"],
            ["react-dom", "19.2.4"],
            ["typescript", "5.x"],
            ["tailwindcss", "4.x"],
            ["@supabase/supabase-js", "2.104.x"],
            ["@supabase/ssr", "0.10.x"],
            ["recharts", "3.8.x"],
            ["lucide-react", "latest stable"],
            ["zod", "3.x"],
        ],
    )
    add_para(doc, "Ver package.json para versiones exactas. npm install ya las congela vía package-lock.json.")

    add_footer(doc, f"{PROYECTO} · Guía de Reconstrucción · {VERSION}")
    out = DOCS_DIR / "06_Guia_Reconstruccion.docx"
    doc.save(out)
    print(f"  ✓ {out.name}")


# ============================================================
# Main
# ============================================================
def main():
    print(f"Generando documentación en {DOCS_DIR}/...")
    print()
    gen_arquitectura()
    gen_diccionario()
    gen_changelog()
    gen_manual()
    gen_guia_ti()
    gen_reconstruccion()
    print()
    print("✅ 6 documentos .docx generados.")
    print(f"   Carpeta: {DOCS_DIR}")


if __name__ == "__main__":
    main()
